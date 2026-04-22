from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Malgun Gothic'
style.font.size = Pt(10.5)

# ── 헬퍼 함수 ──
def h1(text):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(6)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
        run.font.name = 'Malgun Gothic'

def h2(text):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 80, 130)
        run.font.name = 'Malgun Gothic'

def h3(text):
    h = doc.add_heading(text, level=3)
    h.paragraph_format.space_before = Pt(8)
    for run in h.runs:
        run.font.color.rgb = RGBColor(50, 50, 50)
        run.font.name = 'Malgun Gothic'

def body(text, indent=0, space_after=7):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.left_indent = Cm(indent)
    for run in p.runs:
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(10.5)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.name = 'Malgun Gothic'
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.name = 'Malgun Gothic'

def data_box(label, value, sub='', color=(0, 51, 102)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f'▶ {label}: ')
    r1.bold = True
    r1.font.color.rgb = RGBColor(*color)
    r1.font.size = Pt(10.5)
    r1.font.name = 'Malgun Gothic'
    r2 = p.add_run(value)
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.name = 'Malgun Gothic'
    if sub:
        r3 = p.add_run(f'  ({sub})')
        r3.font.size = Pt(9.5)
        r3.font.color.rgb = RGBColor(100, 100, 100)
        r3.font.name = 'Malgun Gothic'

def alert(text, color=(139, 0, 0)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run('※ ' + text)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(*color)
    r.font.name = 'Malgun Gothic'

def source_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run('출처: ' + text)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(120, 120, 120)
    r.font.name = 'Malgun Gothic'
    r.italic = True

def table_multi(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Shading Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Malgun Gothic'
                run.font.bold = True
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = t.rows[i + 1].cells[j]
            cell.text = val
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Malgun Gothic'
                    run.font.size = Pt(10)
    doc.add_paragraph()

def gdp_formula():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run('GDP = C + I + (T − G) + (X − M)')
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0, 51, 102)
    r.font.name = 'Malgun Gothic'

# ══════════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('국내외 금융시장 현황 분석 보고서')
r.bold = True; r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0, 51, 102)
r.font.name = 'Malgun Gothic'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('2026년 4월 21일 기준 | 실증 데이터 기반 심층 분석')
r.font.size = Pt(13); r.font.color.rgb = RGBColor(80, 80, 80)
r.font.name = 'Malgun Gothic'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('전자금융과 디지털화폐')
r.font.size = Pt(15); r.font.color.rgb = RGBColor(0, 51, 102)
r.font.name = 'Malgun Gothic'

doc.add_paragraph()

lines = [
    '■ GDP = C + I + (T−G) + (X−M) 모델 적용',
    '■ GDP · CPI · PPI · 실업률 · 국채발행 실증 데이터',
    '■ 양적완화 · 테이퍼링 · 기준금리 · 인플레이션 이론 적용',
    '■ 트럼프 관세 충격 | 미-이란 전쟁 에너지 위기 | 스태그플레이션 분석',
]
for line in lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.size = Pt(11); r.font.name = 'Malgun Gothic'

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('담당교수: 이중만 교수  |  호서대학교 디지털기술경영학과')
r.font.size = Pt(11); r.font.color.rgb = RGBColor(100, 100, 100)
r.font.name = 'Malgun Gothic'

doc.add_page_break()

# ══════════════════════════════════════════════════
# 서론
# ══════════════════════════════════════════════════
h1('서론')

body(
    '2020년 코로나19 팬데믹 충격에 대응하여 미 연방준비제도(Fed)는 기준금리를 0~0.25%로 인하하고 '
    '매월 1,200억 달러 규모의 국채 및 주택담보부증권(MBS)을 매입하는 양적완화(QE: Quantitative Easing)를 단행하였다. '
    '중앙은행의 본원통화(Monetary Base) 공급이 급증하고 시중은행의 신용창조(Credit Creation)가 맞물리면서 '
    '통화량이 폭발적으로 확대되었다. 그 결과 수요 견인 인플레이션(Demand-Pull Inflation)과 '
    '비용 인상 인플레이션(Cost-Push Inflation)이 동시에 발생하여 2022년 6월 미국 소비자물가지수(CPI)는 '
    '9.1%로 40년 만의 최고치를 기록하였다.'
)
body(
    'Fed는 2022년 3월 테이퍼링(Tapering) 완료 직후 역사상 가장 빠른 금리인상 사이클에 진입하였다. '
    '2022~2023년 기준금리를 0%에서 5.25~5.50%까지 인상하며 인플레이션 억제에 일정 부분 성공하였고, '
    '2024년 9월부터 피벗(Pivot — 통화정책 방향 전환)을 선언하며 금리 인하로 전환하였다. '
    '2025년 12월 Fed는 기준금리를 3.50~3.75%까지 낮추었으나(총 175bp 인하), '
    '2026년 1분기 미국 실질 GDP 성장률이 0.5%로 급락하고 CPI가 3.3%(2026년 3월)로 재상승하는 복합 위기가 전개되고 있다.'
)
body(
    '이 복합 위기의 핵심 원인은 두 가지다. 첫째, 2025년 트럼프 행정부의 전면 관세 부과(실효관세율 7.7% — 1947년 이후 최고)로 '
    '글로벌 무역 구조가 재편되고 비용 인상 인플레이션 압력이 가중되었다. '
    '둘째, 2026년 2월 미-이란 전쟁 발발로 에너지 공급망이 직격탄을 맞으며 '
    'PPI(생산자물가지수) 에너지 부문이 한 달 만에 8.5% 급등하는 등 비용 충격이 확산되고 있다. '
    '본 보고서는 수업에서 학습한 핵심 경제 이론을 바탕으로 실증 데이터를 결합하여 '
    '2026년 4월 현재 국내외 금융시장 현황을 GDP = C + I + (T−G) + (X−M) 모델로 체계적으로 분석한다.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# 제1장. 이론적 배경
# ══════════════════════════════════════════════════
h1('제1장. 이론적 배경 — 수업 핵심 개념 정리')

h2('1.1 화폐금융시스템과 신용창조')
body(
    '현대 금융시스템은 부분지급준비제도(Fractional Reserve Banking)를 근간으로 한다. '
    '은행은 예금액의 10%만을 지급준비금(Reserve Requirement)으로 보유하고 나머지 90%를 대출로 운용하며, '
    '이 과정에서 본원통화가 통화승수(Money Multiplier) 배로 증가하는 신용창조가 이루어진다. '
    '초기 예금 1,000원은 지급준비율 10% 조건에서 최종적으로 1,000원의 신용을 창출하여 총 통화량 1만 원을 형성한다.'
)
body(
    '이 구조는 뱅크런(Bank Run — 예금자들의 동시적 예금 인출)의 위험을 내포한다. '
    '2020년 Fed의 양적완화는 중앙은행이 국채를 대량 매입함으로써 시중에 본원통화를 직접 공급하는 방식이었으며, '
    '이는 신용창조 메커니즘을 통해 통화량을 폭발적으로 확대시키는 주요 인플레이션 원인으로 작용하였다.'
)

h2('1.2 기준금리와 통화정책 전달 메커니즘')
body(
    '기준금리(Base Rate)는 중앙은행이 물가안정과 경제성장의 균형을 위해 결정하는 정책 금리로, '
    '콜금리·국고채금리·COFIX 금리·CD금리 등 시장 금리 전반의 기준이 된다. '
    '금리 조정은 GDP = C + I + (T−G) + (X−M)의 구성요소 전반에 영향을 미친다.'
)
body('금리 인하 경로: 시장 이자율 하락 → 민간소비(C)·기업투자(I) 증가 → GDP 상승. '
     '단, 원화가치 하락 → 수입 물가 상승 → 인플레이션 압력 확대.')
body('금리 인상 경로: 이자율 상승 → 투자(I)·소비(C) 위축 → GDP 하방 압력. '
     '단, 수입물가 안정 → CPI·PPI 하락 → 인플레이션 억제.')

h2('1.3 양적완화(QE)와 테이퍼링(Tapering)')
body(
    '양적완화(QE)는 기준금리가 제로 하한(Zero Lower Bound)에 도달하여 추가 인하가 불가능할 때, '
    '중앙은행이 국채·MBS 등 장기자산을 대량 매입하여 시중 유동성(Liquidity)을 직접 공급하는 비전통적 통화정책이다. '
    '채권 대량 매입 → 채권 수익률(이자율) 하락 → 민간 투자·소비 촉진 → GDP 부양의 경로로 작동한다.'
)
body(
    '테이퍼링(Tapering)은 경기 회복에 따라 QE 규모를 점진적으로 축소하는 출구 전략이다. '
    'Fed는 2021년 11월 테이퍼링을 선언하고 2022년 3월 완료하였다. '
    '테이퍼링 과정에서는 시장의 금리 인상 기대로 자산시장 변동성이 급등하는 "테이퍼 탠트럼(Taper Tantrum)" 현상이 나타난다. '
    '2026년 현재 에너지 충격으로 인플레이션이 재상승하는 국면에서 QE 재개 가능성은 사실상 소멸하였다.'
)

h2('1.4 인플레이션·디플레이션·스태그플레이션')
body('세 가지 물가 국면의 개념과 2026년 현재와의 연결은 다음과 같다.')
bullet('인플레이션(Inflation): 전반적 물가 상승으로 화폐의 구매력이 하락하는 현상. '
       '수요 견인(Demand-Pull Inflation)과 비용 인상(Cost-Push Inflation)으로 구분. '
       '2026년 3월 미국 CPI 3.3%는 관세·에너지 충격이 유발한 비용 인상 인플레이션의 재점화를 반영한다.')
bullet('디플레이션(Deflation): 전반적 물가 하락 현상. 저축의 역설(Paradox of Thrift)로 경기 침체가 심화된다. '
       '일본의 "잃어버린 30년"이 대표 사례. 관세 충격으로 소비가 급격히 위축될 경우 일부 부문에서 디플레이션 압력이 공존할 수 있다.')
bullet('스태그플레이션(Stagflation): 고실업률(경기침체)과 고인플레이션의 동시 발현. '
       '1970년대 오일쇼크가 대표 원인. 2026년 현재 Q1 GDP 성장 0.5%와 CPI 3.3% 재상승이 겹치며 '
       '스태그플레이션 조건이 형성되고 있다. Fed 의장 파월은 "스태그플레이션 위험이 증가했다"고 공식 인정하였다.')
alert('Fed 의장 파월 (2025년): "관세 정책으로 인해 스태그플레이션 위험이 명백히 증가하였다."')
source_note('EY Macroeconomics: Tariff troubles — Could protectionism revive stagflation?')

h2('1.5 GDP 분석 모델')
gdp_formula()
table_multi(
    ['구성요소', '정의', '2026년 4월 현황 요약'],
    [
        ('C — 민간소비', '가계의 재화·서비스 소비 지출 (GDP의 최대 구성요소)', '관세·에너지 가격 상승으로 실질 구매력 압박. 미국 소비심리 위축.'),
        ('I — 투자', '기업 설비·건설 투자 및 재고 변동', '지정학 불확실성으로 기업 투자 관망. 미국 기업투자 +3.2% (AI 주도).'),
        ('(T−G) — 재정수지', '세입(T) − 정부지출(G). 양(+)=재정흑자, 음(−)=재정적자', '한미 모두 대규모 재정적자. 국채 발행 급증. 구축 효과 심화.'),
        ('(X−M) — 순수출', '수출(X) − 수입(M). 국제무역 경쟁력 반영', '관세 보복으로 글로벌 교역량 감소. 미국 순수입 증가, 한국 에너지 수입 비용 급증.'),
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# 제2장. 2026년 글로벌 충격 요인
# ══════════════════════════════════════════════════
h1('제2장. 2026년 글로벌 금융시장 핵심 충격 요인')

h2('2.1 트럼프 관세 충격 — "해방의 날" 이후 (2025~)')
body(
    '2025년 4월, 트럼프 행정부는 "해방의 날(Liberation Day)"이라 명명한 전면 관세를 발표하였다. '
    '모든 수입품에 10% 기본 관세, 중국산에 최고 145%, 캐나다·멕시코에 25%를 부과하였다. '
    '그 결과 미국의 실효 관세율은 7.7%로 상승하였는데, 이는 1947년 이후 최고 수준이다.'
)
table_multi(
    ['항목', '내용', '출처'],
    [
        ('평균 실효관세율', '7.7% (2025년) — 1947년 이후 최고', 'Tax Foundation, 2025'),
        ('미국 가계당 관세 부담', '연간 약 $1,500 (2026년)', 'Penn Wharton Budget Model'),
        ('장기 GDP 영향', '관세 지속 시 장기 GDP −6%, 임금 −5%', 'Penn Wharton Budget Model'),
        ('글로벌 GDP 영향', '세계 실질 GDP −0.5%(2025), −0.7%(2026)', 'IMF 추정'),
        ('기업 부담 비율', '2025년: 기업이 80% 부담 → 2026년: 소비자 전가 비율 급증 전망', 'Axios, 2025'),
    ]
)
body(
    '관세 충격의 GDP 구성요소별 파급 경로는 다음과 같다. '
    '수입품 가격 상승으로 민간소비(C)의 실질 구매력이 감소하고, '
    '글로벌 공급망 재편 불확실성이 기업 투자(I)를 억제하며, '
    '교역 상대국의 보복 관세가 순수출(X−M)을 압박한다. '
    '동시에 수입 물가 상승이 비용 인상 인플레이션(Cost-Push Inflation)을 유발하여 '
    'CPI·PPI가 재상승하고 있다. 2025년에는 기업이 관세의 80%를 흡수하였으나, '
    '2026년에는 소비자에게 전가되는 비율이 급격히 증가할 전망이다.'
)
source_note('Tax Foundation; Penn Wharton Budget Model; CNN Business (2026.01.03)')

h2('2.2 미-이란 전쟁과 에너지 위기 (2026년 2월~)')
body(
    '2026년 2월 28일, 미국·이스라엘 합동 군사작전 "에픽 퓨리(Epic Fury)"가 개시되면서 '
    '중동 에너지 위기가 현실화되었다. 이란 핵시설 및 군사시설에 대한 대규모 공습이 이루어졌으며, '
    '이란은 호르무즈 해협(세계 원유 수송량의 20% 통과)에 기뢰를 설치하고 '
    '하르그 섬(이란 원유수출의 90% 담당) 원유시설이 파괴되었다.'
)
table_multi(
    ['지표', '2026년 3월 데이터', '전쟁 이전 대비 변화'],
    [
        ('국제유가', '급등 (WTI 기준 전쟁 전 대비 40%↑)', '전쟁 이후 4주 연속 상승'),
        ('PPI 에너지 항목', '+8.5% (2026년 3월 MoM)', '2023년 8월 이후 최대 단월 상승'),
        ('해운 컨테이너 할증료', '$3,000 긴급 할증료 부과', '홍해 사태 이후 최고'),
        ('IEA 전략비축유', '사상 최대 4억 배럴 방출 결정', '공급 충격 완화 목적'),
        ('글로벌 증시', '다우존스 400pt+ 급락 (전쟁 직후)', '지정학 위험 프리미엄 급등'),
    ]
)
body(
    '에너지 위기는 2026년 PPI 급등(전년 동기 대비 4.0%, 2026년 3월)의 핵심 동인이다. '
    'PPI는 CPI의 선행 지표로서, PPI 에너지 부문의 급등은 향후 CPI의 추가 상승을 예고한다. '
    '비용 인상 인플레이션(Cost-Push Inflation) 압력이 관세 충격에 더해지면서 '
    '스태그플레이션(Stagflation) 위험이 40년 만에 가장 현실적인 위협으로 부상하고 있다.'
)
source_note('Al Jazeera; IEA; BLS Producer Price Index (2026년 3월); KPMG 경제분석')

doc.add_page_break()

# ══════════════════════════════════════════════════
# 제3장. 미국 금융시장
# ══════════════════════════════════════════════════
h1('제3장. 미국 금융시장 현황 분석 (2026년 4월 기준)')

h2('3.1 GDP 현황 및 구성요소 분석')
body('2026년 1분기(Q1) 미국 실질 GDP 성장률은 전기 대비 연율 0.5%로 집계되었다. '
     '이는 직전 분기(4.4%)에서의 급격한 성장 둔화를 나타내며, 관세 충격과 에너지 위기의 복합 작용 결과다.')

data_box('Q1 2026 실질 GDP 성장률', '0.5% (연율)', '직전 분기 4.4%에서 급락, Ad-hoc News / Atlanta Fed GDPNow')
data_box('민간소비 (C)', 'GDP의 약 68% 차지, 예상 대비 부진', '고금리·에너지 가격 상승으로 내구재 소비 위축')
data_box('기업 투자 (I)', '+3.2% (설비투자 주도)', 'AI 인프라 투자 지속, 주택투자 −2.1% (모기지 금리 6.8%)')
data_box('재정수지 (T−G)', '연방 재정적자 약 $2조+ (GDP 대비 −7%)', '방위비·에너지 지출 급증')
data_box('순수출 (X−M)', '수입 +5.5% vs 수출 +1.2%', '관세 보복에도 국내 수요 강세로 무역적자 확대')
data_box('GDP 물가 디플레이터', '2.8% (Fed 목표 2% 상회)', '인플레이션 압력 지속')

body(
    'GDP 모델 분석: 2026년 Q1 미국 경제에서 민간소비(C)는 에너지 가격 상승과 관세발 소비재 가격 인상으로 '
    '실질 기여도가 하락하고 있다. 투자(I)는 AI 인프라 투자가 일정 부분 지지하고 있으나, '
    '주택투자는 6.8%에 달하는 모기지 금리의 지속으로 위축 상태를 유지하고 있다. '
    '재정수지(T−G)는 방위비 급증으로 역대급 적자를 기록하고 있으며, '
    '이는 대규모 국채 발행 → 시장금리 고착 → 민간 투자(I) 구축 효과(Crowding-out)의 악순환을 심화시킨다. '
    '순수출(X−M)은 관세 보복에도 불구하고 미국의 강한 국내 수요가 수입을 끌어올려 무역적자가 확대되는 구조다.'
)
source_note('BEA GDP Report Q1 2026; Atlanta Fed GDPNow; S&P Global Ratings Economic Outlook Q1 2026')

h2('3.2 소비자물가지수(CPI)와 생산자물가지수(PPI)')

h3('① CPI — 인플레이션의 재점화')
data_box('2022년 6월 CPI (고점)', '9.1%', '40년 만의 최고치 — QE·수요 견인 인플레이션')
data_box('2024년 말 CPI', '2.7%', '금리인상 사이클 효과, 라스트 마일 단계')
data_box('2026년 1월 CPI', '2.4% (YoY)', '피벗 이후 일시 안정')
data_box('2026년 3월 CPI', '3.3% (YoY) / +0.9% (MoM)', '관세·에너지 충격으로 재상승 — BLS', color=(139, 0, 0))
data_box('Core CPI (식품·에너지 제외)', '약 3.5% 수준 유지', '서비스·주거비 "끈적한 인플레이션" 지속')

body(
    '인플레이션 분석: 2022년 인플레이션은 주로 수요 견인(Demand-Pull) — 양적완화 이후 통화량 급증과 '
    '억눌린 소비 수요의 폭발 — 이 원인이었다. 2026년의 CPI 재상승은 성격이 다르다. '
    '관세 부과로 인한 수입품 가격 상승과 에너지 공급 충격이 결합된 비용 인상 인플레이션(Cost-Push Inflation)이 주도한다. '
    '전문 예측기관들은 2026년 헤드라인 CPI가 연율 3.1% 수준을 유지할 것으로 전망한다.'
)
alert('비용 인상 인플레이션은 수요 억제형 금리 인상으로 해소하기 어렵다. 이것이 2026년 Fed 딜레마의 핵심이다.')
source_note('BLS Consumer Price Index Summary (2026년 3월); CNBC CPI breakdown (2026년 1월); Philadelphia Fed SPF Q1 2026')

h3('② PPI — 인플레이션의 선행 경보')
data_box('PPI 2025년 연간', '+3.0% (전년 대비)', '상품 +2.5%, 서비스 +3.2%')
data_box('PPI 2026년 3월', '+4.0% (YoY) / +0.5% (MoM)', '2023년 8월 이후 최고 — BLS')
data_box('PPI 에너지 부문 3월', '+8.5% (MoM)', '에너지 가격 급등 직접 반영 — 역대급 단월 상승')
data_box('PPI 상품 부문 3월', '+1.6% (MoM)', '관세·에너지 복합 충격')

body(
    'PPI는 CPI에 약 2~3개월 선행하는 인플레이션 선행 지표다. '
    '2026년 3월 PPI 에너지 부문의 8.5% 급등과 전체 PPI의 4.0% 상승은 '
    '향후 2분기 CPI가 추가 상승할 가능성이 높음을 강력히 시사한다. '
    '생산자가 원자재·에너지 비용 상승분을 소비자에게 전가하는 과정이 본격화될 것이기 때문이다. '
    '관세 부담도 2025년에는 기업이 80%를 흡수했으나 2026년부터 소비자 전가 비율이 급등할 전망이다.'
)
source_note('BLS Producer Price Index (2026년 3월); BTS Transportation PPI (2026년 3월)')

h2('3.3 실업률 — 노동시장의 균열 신호')
data_box('2024년 초 실업률', '3.7%', '완전고용 수준 유지')
data_box('2026년 2월 실업률', '4.4%', '상승세 뚜렷 — Fed 경고 수준 도달')
data_box('2026년 3월 실업률', '4.3%', '소폭 개선, 실업자 수 720만 명 — BLS')
data_box('2026년 3월 비농업 신규 고용', '+178,000명', '보건·건설·운송 주도, 고점(50만 명) 대비 크게 감소')
data_box('Fed 실업률 전망 (2026년 말)', '4.4%', '성장 둔화 반영 — FOMC 경제전망')

body(
    '2025년 미국 노동시장은 "no-hire, no-fire(신규채용 없이 해고도 없는)" 기조였으나, '
    '2026년에는 관세 충격과 에너지 비용 상승으로 제조업·소매업의 인력 조정이 본격화될 위험이 있다. '
    'Fed 의장 파월은 "일자리 창출이 실제로 마이너스가 될 수 있다"고 경고하였다. '
    '오쿤의 법칙(Okun\'s Law)에 따라 실업률 상승은 민간소비(C)의 추가 위축 압력으로 이어진다.'
)
source_note('BLS Employment Situation (2026년 3월); Yahoo Finance; Federal Reserve FOMC Projections')

h2('3.4 국채 발행 현황 — 재정 위기의 구조화')
data_box('미국 연방 공공부채 총액 (2026년)', '약 $36~37조 (GDP 대비 130% 근접)', 'US Treasury Fiscal Data')
data_box('2026년 1분기 국채 순발행 계획', '$5,780억 (분기 기준)', 'US Treasury Borrowing Advisory')
data_box('2026년 만기 도래 채권', '약 $10조 (전체 부채의 1/3)', '차환 수요 급증 — RIA 분석')
data_box('10년물 국채 수익률', '4.5~5.0% 수준 유지', '대규모 공급 증가로 금리 하락 제한')

body(
    '2026년의 핵심 재정 위험은 구조적 재정적자와 대규모 국채 만기 도래의 결합이다. '
    '전체 미국 국채의 약 1/3($10조)이 2026년에 만기 도래하여 차환 발행이 급증한다. '
    '여기에 미-이란 전쟁으로 인한 방위비 지출 확대가 재정적자를 추가로 확대시키고 있다. '
    '대규모 국채 공급 증가는 채권 수익률(이자율)을 높이는 압력으로 작용하며, '
    '이는 민간 기업의 자금 조달 비용을 높여 투자(I)를 구축(Crowding-out)하는 효과를 유발한다. '
    '수급 관점에서 점점 더 많은 매수자가 필요한 상황이 조성되고 있다.'
)
source_note('RIA (A Third of US Debt Matures in 2026); US Treasury; Charles Schwab Fixed Income Outlook 2026')

h2('3.5 Fed 통화정책 — 극도의 딜레마 국면')
table_multi(
    ['시기', '정책', '기준금리', '핵심 내용'],
    [
        ('2020.03~2021.10', '양적완화(QE)', '0~0.25%', '월 $1,200억 자산 매입, 유동성 공급'),
        ('2021.11~2022.03', '테이퍼링', '0~0.25%', '자산매입 단계적 축소, 3월 완료'),
        ('2022.03~2023.07', '금리인상 사이클', '0%→5.25~5.50%', '빅스텝(0.5%p)·자이언트스텝(0.75%p) 반복'),
        ('2024.09~2025.12', '피벗(금리 인하)', '5.25%→3.50~3.75%', '3회 인하, 총 175bp 인하'),
        ('2026년 (현재)', '동결 기조', '3.50~3.75%', '2026년 1회 추가 인하만 점도표 반영, 관망'),
    ]
)
body(
    'Fed의 2026년 딜레마: 2025년 12월 FOMC 점도표(Dot Plot)는 2026년 1회 금리 인하만을 시사하였다. '
    'CPI가 3.3%로 재상승하고 PPI 에너지 부문이 8.5% 급등하는 상황에서 금리를 내리면 인플레이션 재가속 위험이 있고, '
    '금리를 올리면 이미 0.5%로 급락한 GDP 성장이 마이너스로 진입할 수 있다. '
    'Fed는 2026년 2분기 현재 에너지 충격의 일시성 여부와 관세 파급 효과를 관망하며 '
    '동결 기조를 유지하고 있다. 이는 공급 충격 앞에서 전통적 통화정책의 한계를 드러내는 전형적인 스태그플레이션 딜레마다.'
)
source_note('Federal Reserve FOMC Statement (2025년 12월); Fed Dot Plot (2025년 12월); JP Morgan Global Research')

doc.add_page_break()

# ══════════════════════════════════════════════════
# 제4장. 한국 금융시장
# ══════════════════════════════════════════════════
h1('제4장. 한국 금융시장 현황 분석 (2026년 4월 기준)')

h2('4.1 GDP 현황 및 구성요소 분석')
data_box('2025년 실질 GDP 성장률', '약 0.7~1.0%', 'HRI 하향조정 0.7% / 한국은행 1.0% 추정')
data_box('2026년 GDP 성장률 전망', '약 1.9~2.1%', 'KDI·OECD 전망, 반도체 수출 회복 기대')

body(
    '2025년 한국 경제는 트럼프 관세 충격과 글로벌 수요 위축으로 당초 전망을 크게 하회하였다. '
    '현대경제연구원(HRI)은 관세 충격 이후 2025년 성장률을 0.7%로 대폭 하향 조정하였으며, '
    '이는 코로나19 충격을 제외하면 2009년 금융위기 이후 최저 수준이다.'
)
table_multi(
    ['구성요소', '2025년 현황', '2026년 전망', '주요 요인'],
    [
        ('C — 민간소비', '전년 대비 +약 1.0%', '에너지 가격 급등으로 추가 둔화', '가계 실질소득 감소, 에너지·식품 가격 상승'),
        ('I — 투자', '반도체 회복, 건설 위축', 'AI 반도체 수요 지속, 기업투자 불확실', '글로벌 AI 투자 수요 + PF 부실 리스크'),
        ('(T−G) 재정수지', '관리재정적자 약 −100조 원', '방위비·에너지 보조금 확대', 'GDP 대비 −4.5% 내외'),
        ('(X−M) 순수출', '약 +900억 달러 흑자', '에너지 수입 급증으로 흑자 축소', '반도체 수출 vs. 에너지 수입 비용 대비'),
    ]
)
body(
    'GDP 모델 분석: 한국의 GDP 모델에서 2026년 가장 큰 위험 요인은 재정수지(T−G)의 추가 악화와 '
    '순수출(X−M) 흑자의 감소다. '
    '에너지 수입 의존도가 높은 한국은 유가 40% 급등이 수입비용을 폭발적으로 증가시켜 '
    '(X−M) 흑자를 잠식한다. 반도체 수출이 일정 부분 이를 상쇄하고 있으나, '
    '글로벌 경기 둔화가 반도체 수요를 약화시킬 경우 (X−M) 구성요소가 GDP에 음(−)의 기여를 할 위험이 있다.'
)
source_note('KDI 경제전망 (2026년 2월); OECD Economic Outlook; 현대경제연구원 (HRI, 2025년 4월)')

h2('4.2 CPI와 PPI')
data_box('2024년 CPI', '약 2.3%', '한국은행 목표치(2%) 근접, 안정 국면')
data_box('2025년 CPI', '약 2.0%', '물가 목표 달성, 안정세 유지')
data_box('2026년 CPI 전망', '약 2.1~2.2%', '한국은행 전망 — 에너지 충격 반영 시 상향 위험')
data_box('2026년 PPI', '에너지 부문 중심 상승 압력', '국제유가 40% 급등의 직접 파급', color=(180, 80, 0))

body(
    '한국 CPI는 2024~2025년 안정세를 유지하였으나, 2026년 에너지 충격이 새로운 상승 압력으로 작용하고 있다. '
    '국제유가 급등은 경유·도시가스·전기요금을 통해 국내 물가 전반으로 전이된다. '
    '특히 PPI 에너지 부문의 선행적 상승이 향후 CPI 재상승을 예고하는 점이 우려된다. '
    '수요 측 인플레이션 압력은 약하나, 에너지발 비용 인상 인플레이션이 2026년 한국 물가의 핵심 위험이다.'
)
source_note('한국은행 경제통계시스템(ECOS); 통계청 소비자물가지수')

h2('4.3 실업률')
data_box('2025~2026년 실업률', '약 2.8%', '사실상 완전고용 유지 — Trading Economics')
data_box('청년 실업률', '약 6.0~7.0%', '전체 실업률의 2배 이상, 구조적 취약 계층')
data_box('고용 구조 우려', '제조업 고용 조정 위험', '에너지 비용 급증으로 수익성 악화 업종 중심')

body(
    '한국 실업률은 2.8%로 공식 완전고용 수준을 유지하고 있으나, '
    '이는 비경제활동인구 증가와 플랫폼 노동 확산으로 공식 통계에서 포착되지 않는 '
    '"체감 청년 실업"이 은폐된 결과일 수 있다. '
    '에너지 비용 급등으로 원가 부담이 가중되는 에너지 집약적 제조업 부문에서 '
    '2026년 하반기 고용 조정이 발생할 경우 실업률 상승이 가시화될 수 있다.'
)
source_note('통계청 고용동향; Trading Economics 한국 실업률')

h2('4.4 국채 발행 현황')
data_box('2025년 국고채 발행', '약 168조 원', '재정적자 보전 + 만기 차환')
data_box('2026년 국고채 발행 전망', '168조 원 초과 예상', '에너지 보조금·방위비 지출 확대')
data_box('국고채 10년물 금리', '3.0~3.5% 수준', '한국은행 금리 인하 기조에도 발행 증가로 제한')
data_box('국가채무비율', 'GDP 대비 약 50%', '선진국 대비 낮으나 빠른 증가 속도가 우려')

h2('4.5 한국은행 통화정책')
data_box('2024년 10월 기준', '금리인하 사이클 시작', '3.50%에서 인하 개시')
data_box('2025년 5월 이후', '기준금리 2.50%', '총 100bp 인하 후 동결 — 6회 연속 동결 (2026년 2월 포함)')
data_box('2026년 GDP 전망 (한국은행)', '2.0%', '기존 1.8%에서 상향, 반도체 수출 회복 반영')
data_box('2026년 CPI 전망 (한국은행)', '2.2%', '에너지 충격 반영 전 전망치')

body(
    '한국은행은 2026년 2월 기준금리를 2.5%로 동결(만장일치)하였다. '
    '추가 인하를 제약하는 요인은 두 가지다. '
    '첫째, 에너지 충격으로 CPI 재상승 위험이 높아진 상황에서 금리 인하는 인플레이션을 자극할 수 있다. '
    '둘째, 금리 인하는 원화 약세를 유발하여 에너지 수입 비용을 추가로 높이는 역효과를 낳는다. '
    '반면 동결·인상은 이미 위축된 내수를 추가로 압박한다. '
    '한국은행 역시 Fed와 마찬가지로 스태그플레이션 딜레마에 직면해 있다.'
)
source_note('한국은행 통화정책방향 (2026년 2월); Korea Biz Review; Toss Bank 기준금리 분석')

doc.add_page_break()

# ══════════════════════════════════════════════════
# 제5장. GDP 모델 한미 비교
# ══════════════════════════════════════════════════
h1('제5장. GDP = C + I + (T−G) + (X−M) 모델 적용 — 한미 비교 분석')

gdp_formula()

table_multi(
    ['구성요소', '미국 (2026년 Q1)', '한국 (2026년 추정)', '시사점'],
    [
        ('GDP 성장률', '+0.5% (연율)', '+1.0~1.5% (추정)', '미국의 충격이 더 크게 나타남'),
        ('C — 민간소비', 'GDP의 68%, 에너지 압박으로 둔화', 'GDP의 48%, 에너지·물가로 실질소득 감소', '양국 소비 모두 위축 압력'),
        ('I — 투자', '+3.2% (AI 주도)', '반도체 지속, 건설 부진', '미국 AI 투자가 버팀목'),
        ('(T−G) 재정수지', '−$2조+ (GDP −7%)', '−100조 원+ (GDP −4.5%)', '양국 모두 대규모 적자, 국채 급증'),
        ('(X−M) 순수출', '무역적자 확대 (수입 +5.5%)', '흑자 유지, 에너지 수입 급증으로 축소', '한국 에너지 취약성 노출'),
        ('CPI', '3.3% (3월, 재상승)', '2.1~2.2% (에너지 반영 전)', '미국 충격 선행, 한국 후행 가능성'),
        ('PPI', '4.0% YoY, 에너지 +8.5% MoM', '에너지 중심 상승 압력', 'PPI → CPI 선행 전이 예고'),
        ('실업률', '4.3% (상승 추세)', '2.8% (안정)', '한국 노동시장 상대적 양호'),
        ('기준금리', '3.50~3.75% (동결)', '2.50% (동결)', '양국 모두 딜레마로 동결 기조'),
    ]
)

h2('5.1 핵심 분석 — 스태그플레이션 조건의 형성')
body(
    '2026년 4월 현재, 미국과 한국 모두 GDP 모델의 4개 구성요소가 동시에 압박받는 '
    '"사면초가(四面楚歌)" 국면에 처해 있다.'
)
bullet('C (민간소비) 위축: 관세로 인한 소비재 가격 상승 + 에너지 비용 증가로 가계 실질 구매력 감소.')
bullet('I (투자) 위축: 지정학 불확실성 + 높은 시장 금리가 기업 투자 의사결정을 억제.')
bullet('(T−G) 대규모 음수: 재정적자 확대 → 국채 발행 급증 → 시장 금리 고착 → 민간 투자(I) 구축 효과. 악순환.')
bullet('(X−M) 악화 압력: 미국은 관세 보복으로 수출 감소, 한국은 에너지 수입비용 폭증으로 흑자 잠식.')

body(
    '스태그플레이션(Stagflation)의 두 조건 — 높은 인플레이션(미국 CPI 3.3%, PPI 4.0%)과 '
    '경기침체(미국 Q1 GDP 0.5%) — 이 동시에 충족되고 있다. '
    '1970년대 오일쇼크와 유사하게, 이번에도 에너지 공급 충격(Cost-Push)이 핵심 동인이다. '
    '그러나 현 국면은 관세발 인플레이션이 추가로 겹쳐 있어 1970년대보다 정책 대응이 더 복잡하다.'
)

h2('5.2 통화정책과 GDP 구성요소 — 딜레마 매트릭스')
table_multi(
    ['정책 방향', '인플레이션 효과', 'GDP 효과', '순효과 평가'],
    [
        ('금리 인하', 'CPI 재가속 위험 (부정)', 'C·I 지지 (긍정)', '인플레이션 악화로 불가'),
        ('금리 동결', 'CPI 억제 일부 (중립)', '성장 회복 지연 (부정)', '현재 선택 — 관망 기조'),
        ('금리 인상', 'CPI 억제 (긍정)', 'GDP 마이너스 위험 (강한 부정)', '경기침체 악화로 불가'),
    ]
)
alert('스태그플레이션 환경에서는 어떤 금리 정책도 부작용을 수반한다. 공급 측 해결(에너지 안보, 관세 협상)이 근본 처방이다.')

doc.add_page_break()

# ══════════════════════════════════════════════════
# 제6장. 종합 평가 및 시사점
# ══════════════════════════════════════════════════
h1('제6장. 종합 평가 및 시사점')

h2('6.1 2026년 국내외 금융시장의 구조적 특징 — 4가지')

h3('① 스태그플레이션의 현실화')
body(
    '미국 Q1 2026 GDP 성장률 0.5%와 CPI 3.3%의 동시 발현은 '
    '경기침체와 인플레이션의 동시 진행이라는 스태그플레이션의 조건을 충족하고 있다. '
    'Fed 의장 파월이 공식적으로 스태그플레이션 위험을 인정한 것은 시장에 강력한 경고 신호다. '
    '에너지 공급 충격이 단기간에 해소되지 않을 경우, 스태그플레이션은 구조화될 위험이 있다.'
)

h3('② 비용 인상 인플레이션의 이중 구조')
body(
    '2026년 인플레이션은 2022년과 달리 수요 견인(Demand-Pull)이 아닌 '
    '비용 인상(Cost-Push) 중심의 이중 충격 구조다. '
    '관세 부과로 인한 수입품 가격 상승(관세발 비용 인상)과 에너지 공급 충격(에너지발 비용 인상)이 '
    'PPI를 4.0%로 밀어올리고, 이것이 2~3개월 시차를 두고 CPI를 추가 상승시키는 '
    '"인플레이션 파이프라인"이 가동 중이다.'
)

h3('③ 재정의 구조적 악화와 국채 시장 위험')
body(
    '미국 연방부채의 1/3($10조)이 2026년 만기 도래하는 상황에서 방위비 급증이 재정적자를 추가 확대하고 있다. '
    '분기별 국채 순발행이 $5,780억에 달하며, 이는 국채 수익률(10년물 4.5~5.0%)의 고착과 '
    '민간 투자(I) 구축 효과를 유발하고 있다. '
    'GDP 모델에서 (T−G)의 대규모 음수 확대는 단기 수요를 지지하는 효과보다 '
    '장기적으로 이자율 상승을 통해 C와 I를 억제하는 부작용이 더 클 수 있다.'
)

h3('④ 한국의 에너지 구조적 취약성')
body(
    '에너지 자급률이 낮은 한국은 국제유가 40% 급등의 직접적 피해국이다. '
    '에너지 수입비용 급증은 GDP 모델에서 순수출(X−M) 흑자를 잠식하고, '
    'PPI를 통해 CPI 상승 압력을 전이하며, '
    '원화 약세를 유발하여 에너지 수입비용을 추가 증가시키는 악순환을 형성한다. '
    '반도체 수출 경쟁력이 유지되더라도 에너지 구조적 취약성은 '
    '한국 경제의 외부 충격 흡수 능력을 근본적으로 제약한다.'
)

h2('6.2 GDP 모델 구성요소별 정책 시사점')
table_multi(
    ['구성요소', '현황 진단', '정책 방향'],
    [
        ('C — 민간소비', '에너지·관세 압박으로 실질 구매력 감소', '에너지 바우처, 취약계층 소비 지원. 관세 협상으로 수입품 가격 인하.'),
        ('I — 투자', '불확실성으로 기업투자 관망', 'AI·반도체·재생에너지 분야 투자 인센티브. 지정학 위험 관리.'),
        ('(T−G) 재정수지', '방위비·보조금으로 적자 확대', '지출 우선순위 조정. 재정 준칙 강화로 채권시장 신뢰 유지.'),
        ('(X−M) 순수출', '교역량 감소, 에너지 수입비용 급증', '관세 협상·무역 다변화. 에너지 수입원 다변화 및 재생에너지 전환 가속화.'),
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# 결론
# ══════════════════════════════════════════════════
h1('결론')

body(
    '본 보고서는 호서대학교 "전자금융과 디지털화폐" 수업에서 학습한 핵심 경제 개념 — '
    '화폐금융시스템, 신용창조, 기준금리, 양적완화(QE), 테이퍼링, 인플레이션·디플레이션·스태그플레이션 — 을 '
    '이론적 기반으로 삼아, 실증 데이터를 결합하여 2026년 4월 현재 국내외 금융시장을 '
    'GDP = C + I + (T−G) + (X−M) 모델로 분석하였다.'
)
body(
    '실증 데이터가 보여주는 2026년 금융시장의 실상은 다음과 같다. '
    '미국 Q1 GDP 성장률은 0.5%로 급락하였고(BEA/Atlanta Fed), '
    'CPI는 3.3%로 재상승하였으며(BLS, 2026년 3월), '
    'PPI 에너지 부문은 월 8.5% 급등하였고(BLS), '
    '실업률은 4.3%로 상승 추세를 보이고 있다(BLS). '
    'Fed는 기준금리를 3.50~3.75%에 동결하며 2026년 1회 인하만을 시사하는 극도의 관망 기조를 취하고 있다.'
)
body(
    '한국 경제는 2025년 GDP 성장률 0.7~1.0%로 충격을 받았으며, '
    '한국은행은 기준금리를 2.50%에 동결(6회 연속)하며 에너지 충격과 성장 둔화 사이의 딜레마를 관리하고 있다. '
    '에너지 의존형 수입 구조가 순수출(X−M) 흑자를 잠식하는 것이 2026년 한국 경제의 핵심 취약 고리다.'
)
body(
    '이론과 실증이 일치하는 결론은 명확하다. '
    '2026년 현재 국내외 금융시장은 1970년대 오일쇼크 이후 가장 심각한 스태그플레이션 위험에 직면해 있으며, '
    '전통적 통화정책(기준금리 조정)만으로는 에너지 공급 충격과 관세발 비용 인상 인플레이션을 동시에 해결할 수 없다. '
    '에너지 안보의 외교·군사적 해결, 관세 협상을 통한 무역 정상화, '
    '재정 건전성 관리라는 공급 측 처방이 통화정책과 병행되어야 한다.'
)
body(
    'GDP = C + I + (T−G) + (X−M) 모델은 단순한 공식이 아니다. '
    '이 모델의 4개 구성요소를 통해 우리는 소비자의 지갑 사정, 기업의 투자 심리, '
    '정부의 재정 건전성, 그리고 국가 간 교역 경쟁력이라는 경제의 4개 기둥이 '
    '동시에 흔들리는 2026년의 복합 위기를 총체적으로 진단할 수 있다.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════
# 참고자료
# ══════════════════════════════════════════════════
h1('참고자료')

h2('강의 자료')
refs_lecture = [
    '이중만 (2022). 전자금융과 디지털화폐 — 화폐 금융시스템. 호서대학교 디지털기술경영학과.',
    '장세형 (2019). 블록체인 & 암호화폐 105문답.',
    '임명환 (2020). 블록체인의 개요 및 융복합.',
]
for r in refs_lecture:
    body(r, indent=0.5, space_after=3)

h2('미국 정부·연구기관')
refs_us = [
    'U.S. Bureau of Economic Analysis (BEA). (2026). Q1 2026 GDP Advance Estimate.',
    'U.S. Bureau of Labor Statistics (BLS). (2026). Consumer Price Index Summary — March 2026.',
    'U.S. Bureau of Labor Statistics (BLS). (2026). Producer Price Index Summary — March 2026.',
    'U.S. Bureau of Labor Statistics (BLS). (2026). Employment Situation — March 2026.',
    'Federal Reserve. (2025). FOMC Statement and Economic Projections — December 2025.',
    'U.S. Department of the Treasury. (2026). Treasury Borrowing Advisory Committee.',
    'Federal Reserve Bank of Atlanta. (2026). GDPNow Real-Time Estimate.',
    'Philadelphia Fed. (2026). Survey of Professional Forecasters Q1 2026.',
]
for r in refs_us:
    body(r, indent=0.5, space_after=3)

h2('국제기관·연구소')
refs_intl = [
    'International Monetary Fund (IMF). (2025~2026). World Economic Outlook.',
    'Penn Wharton Budget Model. (2025). Economic Effects of President Trump\'s Tariffs.',
    'Tax Foundation. (2025). Tracking the Impact of the Trump Tariffs & Trade War.',
    'S&P Global Ratings. (2026). Economic Outlook US Q1 2026.',
    'Deloitte Insights. (2026). US Economic Forecast Q1 2026.',
    'International Energy Agency (IEA). (2026). Strategic Petroleum Reserve Release.',
    'KPMG Korea. (2025). 2026년 경제 및 산업 전망.',
]
for r in refs_intl:
    body(r, indent=0.5, space_after=3)

h2('한국 정부·연구기관')
refs_kr = [
    '한국은행. (2026). 통화정책방향 — 2026년 2월.',
    '한국은행. (2025). 통화신용정책보고서 2025년 9월.',
    '한국개발연구원(KDI). (2026). 경제전망 수정 (2026년 2월).',
    '현대경제연구원(HRI). (2025). 2025년 한국 경제 전망 수정.',
    '기획재정부. (2025~2026). 국채 발행 및 관리재정수지.',
    '통계청. (2025~2026). 국민계정·고용동향.',
]
for r in refs_kr:
    body(r, indent=0.5, space_after=3)

h2('미디어·분석 리포트')
refs_media = [
    'EY US. (2025). Tariff troubles: Could protectionism revive stagflation?',
    'CNN Business. (2026.01.03). Tariffs could really sting in 2026.',
    'Axios. (2025.06.03). Trump tariffs: Economic forecast shows short-lived inflationary impact.',
    'JP Morgan Global Research. (2026). US Tariffs: What\'s the Impact?',
    'Time / Davos. (2026). Why Trump\'s Tariffs Are Like Termites.',
    'Yahoo Finance. (2026). The US labor market ground to a halt in 2025.',
    'CNBC. (2026.01.13). Here\'s the inflation breakdown for December 2025.',
    'Al Jazeera. (2026). Iran war day 14 updates.',
]
for r in refs_media:
    body(r, indent=0.5, space_after=3)

# 저장
output_path = '/Users/kylechoi/Desktop/Ai_works/output/금융시장분석/2026-04-21_국내외금융시장분석/국내외_금융시장_현황분석_심층_2026.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'Word 파일 생성 완료: {output_path}')
