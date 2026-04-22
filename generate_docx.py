from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# 페이지 여백 설정
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Malgun Gothic'
font.size = Pt(10.5)

# ── 표지 ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n\n')
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('미국-이란 전쟁 뉴스 조사 보고서')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('(2026년 3월 14일 기준)')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(80, 80, 80)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('설교 예화 자료')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('"복음은 시선을 돌린다 — 바른 시선을 보게 한다"')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(139, 0, 0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n\n\n작성일: 2026년 3월 14일')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ── 목차 ──
h = doc.add_heading('목차', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

toc_items = [
    '1. 전쟁 개요 (타임라인)',
    '2. 피해 현황 (14일차 기준)',
    '3. 군사적 상황',
    '4. 경제적 영향',
    '5. 국제사회 반응',
    '6. 휴전 협상 현황',
    '7. 설교 예화를 위한 데이터 분석 포인트',
    '8. 출처 (Sources)',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    for run in p.runs:
        run.font.size = Pt(12)

doc.add_page_break()

# ── 헬퍼 함수 ──
def add_section_heading(text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_sub_heading(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 80, 130)

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════
# 1. 전쟁 개요
# ══════════════════════════════════════════════════
add_section_heading('1. 전쟁 개요 (타임라인)')

add_sub_heading('배경')
add_bullet('2025년 8월: 프랑스·독일·영국이 UN 핵합의 \'스냅백\' 제재 발동')
add_bullet('2025년 9월 27일: UN 대이란 제재 자동 복원 → 이란 리알화 폭락, 12월 대규모 시위')
add_bullet('2025~2026년 초: 미국-이란 핵 협상 진행 (제네바), 이란은 타협 의사 표명')
add_bullet('2026년 2월 25일: 미국, 핵 협상 직전 새로운 대이란 제재 발표')

add_sub_heading('전쟁 발발 (2026년 2월 28일, Day 1)')
add_bullet('미국-이스라엘 합동 군사작전 \'에픽 퓨리(Epic Fury)\' 개시')
add_bullet('이란 최고지도자 알리 하메네이 암살 (3월 1일 사망 공식 발표)')
add_bullet('이란의 핵시설, 탄도미사일 기지, 군사시설 대규모 공습')

add_sub_heading('전쟁 확산 (Day 2~7)')
add_bullet('이란 즉각 보복: 500발 이상의 탄도·해상 미사일, 2,000대 드론 발사 (40% 이스라엘 향, 60% 중동 내 미군기지 향)')
add_bullet('헤즈볼라, 이스라엘에 미사일·드론 공격 재개 → 2026 레바논 전쟁으로 확전')
add_bullet('두바이 제벨알리 항구, 이란 미사일 잔해로 화재 (3월 1일)')
add_bullet('호르무즈 해협 선박 공격 시작')

add_sub_heading('전쟁 격화 (Day 8~14, 3월 7일~13일)')
add_bullet('3월 8일: 모즈타바 하메네이, 아버지 후임 최고지도자로 선출')
add_bullet('3월 10일: 미국 "역대 가장 강력한 이란 내 공습" 시행')
add_bullet('3월 10일: 트럼프 "전쟁 곧 끝날 것" 발언, 러시아 중재 정상통화')
add_bullet('3월 11일: 이란, 호르무즈 해협에 기뢰 설치 시작')
add_bullet('3월 13일: 미국, 하르그 섬(이란 원유수출 90%) "모든 군사 목표 완전 파괴" 선언')
add_bullet('3월 13일: 이란, 트럼프 측 휴전 제의 2차례 거부')

doc.add_page_break()

# ══════════════════════════════════════════════════
# 2. 피해 현황
# ══════════════════════════════════════════════════
add_section_heading('2. 피해 현황 (14일차 기준)')

add_sub_heading('인명 피해')
# 표로 정리
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = '국가/세력'
hdr[1].text = '사망'
hdr[2].text = '부상'
data = [
    ('이란', '1,444명 이상 (아동 약 200명)', '18,551명'),
    ('레바논', '570명 이상', '-'),
    ('이스라엘', '12명', '-'),
    ('미군', '7명', '중상 8명'),
]
for i, (a, b, c) in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = a
    row[1].text = b
    row[2].text = c

doc.add_paragraph()

add_sub_heading('민간 기반시설 피해 (이란적십자 발표)')
add_bullet('민간 건물 약 20,000동 피해 (주거시설 16,000동 이상)')
add_bullet('의료시설 77곳 피해')
add_bullet('학교 65곳 피해')
add_bullet('2월 28일 민바브(호르모즈간 주) 학교 폭격: 아동 150명 포함 165명 사망 → UN "인도법의 중대한 위반"으로 규정')

add_sub_heading('난민/실향민')
add_bullet('320만 명 이상 이란 내 실향민 발생 (2주 만에)')
add_bullet('대부분 테헤란·남부 도시에서 북부·농촌 지역으로 피난')
add_bullet('레바논: 30만 명 이상 집을 잃고 대피소 포화 상태')

doc.add_page_break()

# ══════════════════════════════════════════════════
# 3. 군사적 상황
# ══════════════════════════════════════════════════
add_section_heading('3. 군사적 상황')

add_sub_heading('미국-이스라엘 측')
add_bullet('이란 미사일 역량 90% 감소, 드론 역량 95% 감소 (미 국방장관 발표)')
add_bullet('하르그 섬 원유시설 완전 파괴 (이란 원유수출의 90% 담당)')
add_bullet('핵시설 및 주요 군사시설 집중 타격')

add_sub_heading('이란 측')
add_bullet('호르무즈 해협 기뢰 설치 (세계 원유 수송량의 20% 통과)')
add_bullet('"미군 기지 폐쇄 않으면 공격 계속" (모즈타바 하메네이)')
add_bullet('"장기전으로 미국에 고통을 주겠다" — 결사항전 선언')
add_bullet('중동 내 협력국 석유시설 보복 공격 경고')

add_sub_heading('헤즈볼라/레바논')
add_bullet('이스라엘에 대한 "존재론적 전투" 선언')
add_bullet('이스라엘, 레바논 내 공습 대폭 확대')

# ══════════════════════════════════════════════════
# 4. 경제적 영향
# ══════════════════════════════════════════════════
add_section_heading('4. 경제적 영향')

add_sub_heading('원유/에너지')
add_bullet('국제유가 전쟁 이후 40% 이상 폭등')
add_bullet('호르무즈 해협 선박 운항 사실상 중단')
add_bullet('국제에너지기구(IEA): 사상 최대 4억 배럴 전략비축유 방출 결정')
add_bullet('카타르 에너지장관: "전쟁 계속되면 걸프 산유국 수출 중단 불가피 → 세계 경제 붕괴 우려"')

add_sub_heading('금융시장')
add_bullet('다우존스 3월 2일 400포인트 이상 급락')
add_bullet('세계 경기침체 진입 경고')
add_bullet('해운 컨테이너 긴급할증료 $3,000 부과')

add_sub_heading('인도주의 위기 파급')
add_bullet('글로벌 구호물자 공급망 마비')
add_bullet('기존 인도주의 위기 지역(아프리카, 아시아 등)에도 연쇄 타격')

doc.add_page_break()

# ══════════════════════════════════════════════════
# 5. 국제사회 반응
# ══════════════════════════════════════════════════
add_section_heading('5. 국제사회 반응')

add_sub_heading('지지/협력')
add_bullet('프랑스: 미군에 프랑스 기지 사용 허가 (3월 5일), 항공모함 지중해 파견')
add_bullet('영국: 방어 목적 영국 군사기지 사용 허가')
add_bullet('포르투갈: 라제스 비행장 방어 목적 사용 허가')
add_bullet('캐나다: "참전 가능성 배제 못 한다"')

add_sub_heading('중립/평화 호소')
add_bullet('남아공, 아제르바이잔, 인도네시아, 레바논, 아일랜드, 슬로베니아, 바티칸, 우루과이, 우즈베키스탄, 베네수엘라 — 평화적 해결 촉구')

add_sub_heading('국제기구')
add_bullet('UN 인권최고대표사무소(OHCHR): "최악의 우려 현실화, 대화만이 유일한 출구"')
add_bullet('국제앰네스티: 민간인 보호 및 국제법 준수 긴급 촉구')
add_bullet('WHO: 이란 내 아동 사망 200명 보고')

add_sub_heading('러시아')
add_bullet('3월 10일 정상통화를 통한 중재 시도')

# ══════════════════════════════════════════════════
# 6. 휴전 협상 현황
# ══════════════════════════════════════════════════
add_section_heading('6. 휴전 협상 현황')

add_sub_heading('미국 측')
add_bullet('트럼프: "이란 지도부와 대화할 의향 있다", "전쟁은 곧 끝날 것"')
add_bullet('"더 많은 것을 해야 한다(more of the same)" — 군사적 압박 병행')

add_sub_heading('이란 측')
add_bullet('이란 외무장관: "영구적 전쟁 종료만 수용"')
add_bullet('휴전 선결조건: (1) 공습 재발 방지 국제적 확약 (2) 배상금 지급')
add_bullet('트럼프 측 휴전 제의 2차례 거부 (3월 12일)')

add_sub_heading('전망')
add_bullet('미국 언론 분석 5개 시나리오 — 상황은 여전히 불투명하며, 양측 모두 "양보 없는 승리"를 추구 중')

doc.add_page_break()

# ══════════════════════════════════════════════════
# 7. 설교 예화 분석 포인트
# ══════════════════════════════════════════════════
add_section_heading('7. 설교 예화를 위한 데이터 분석 포인트')

p = doc.add_paragraph()
run = p.add_run('"복음은 시선을 돌린다: 바른 시선을 보게 한다"')
run.bold = True
run.italic = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(139, 0, 0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 분석 포인트 1
add_sub_heading('[분석 포인트 1] 세상이 보는 시선 — 두려움과 불안')
add_bullet('전 세계가 유가 폭등, 경기침체, 군사적 확전에 시선이 고정되어 있다.')
add_bullet('주식시장 폭락, 에너지 위기, 난민 위기 — 모든 뉴스가 "두려움"을 전한다.')
add_bullet('사람들의 시선은 \'생존\'과 \'안보\'에 향해 있다.')

# 분석 포인트 2
add_sub_heading('[분석 포인트 2] 가려진 시선 — 보이지 않는 사람들')
add_bullet('150명의 학교 아이들이 폭격으로 사망했으나, 세계의 관심은 유가와 경제에 집중.')
add_bullet('320만 명의 실향민은 숫자로만 존재하고, 그들의 이름과 얼굴은 보이지 않는다.')
add_bullet('기존 인도주의 위기 지역(아프리카, 아시아)은 공급망 마비로 더 큰 고통을 받지만 뉴스의 조명을 받지 못한다.')

# 분석 포인트 3
add_sub_heading('[분석 포인트 3] 각자의 시선 — 같은 전쟁, 다른 프레임')
add_bullet('', bold_prefix='미국: ')
p = doc.paragraphs[-1]
p.clear()
run = p.add_run('미국: ')
run.bold = True
run.font.size = Pt(10.5)
run = p.add_run('"이란의 핵위협 제거, 국가 안보" → 안보의 시선')
run.font.size = Pt(10.5)

add_bullet('"침략에 대한 정당한 저항, 결사항전" → 생존의 시선', bold_prefix='이란: ')
add_bullet('"세계 경제 위기, 에너지 안보" → 이익의 시선', bold_prefix='국제사회: ')
add_bullet('"민간인 보호, 대화와 평화" → 생명의 시선', bold_prefix='바티칸/평화단체: ')

p = doc.add_paragraph()
run = p.add_run('→ 모두가 "자기 쪽에서" 세상을 바라보고 있다.')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(139, 0, 0)

# 분석 포인트 4
add_sub_heading('[분석 포인트 4] 복음이 돌리는 시선 — 예화 적용 방향')
add_bullet('세상은 "누가 이기느냐"에 시선을 두지만, 복음은 "누가 아파하느냐"에 시선을 돌린다.')
add_bullet('세상은 유가 차트와 주가 그래프에 시선을 두지만, 복음은 폭격 속 학교 아이들의 얼굴에 시선을 돌린다.')
add_bullet('세상은 "어떻게 하면 내가 살아남느냐"에 시선을 두지만, 복음은 "어떻게 하면 저 사람이 살 수 있느냐"에 시선을 돌린다.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('성경적 근거:')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0, 51, 102)

add_bullet('예수님은 군중이 "빵"을 보고 몰려올 때, "영혼"을 보셨다. (요 6:26-27)')
add_bullet('제자들이 "폭풍"을 볼 때, 예수님은 "믿음"을 보셨다. (막 4:40)')
add_bullet('사람들이 "세리와 죄인"을 볼 때, 예수님은 "하나님의 자녀"를 보셨다. (눅 15)')

# 분석 포인트 5
add_sub_heading('[분석 포인트 5] 구체적 예화 소재')

quotes = [
    '"2주 만에 320만 명이 집을 잃었습니다. 그런데 그날 전 세계 뉴스의 헤드라인은 \'유가 40% 폭등\'이었습니다. 우리의 시선은 어디에 있었습니까?"',
    '"150명의 아이들이 학교에서 폭격으로 죽었습니다. 같은 날, 세계 증시는 \'에너지 안보\'를 논했습니다. 복음은 우리에게 묻습니다 — 당신은 차트를 보고 있습니까, 아이들의 얼굴을 보고 있습니까?"',
    '"카타르 에너지장관은 \'세계 경제가 무너질 수 있다\'고 경고했습니다. 그러나 바티칸은 \'한 사람의 생명이 세계 경제보다 귀하다\'고 말했습니다. 복음은 시선의 방향을 바꿉니다."',
]
for q in quotes:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(q)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# ══════════════════════════════════════════════════
# 8. 출처
# ══════════════════════════════════════════════════
add_section_heading('8. 출처 (Sources)')

add_sub_heading('영문 출처')
en_sources = [
    'Al Jazeera: Iran war Day 14 updates',
    'CNN: What we know on the 14th day of the US and Israel\'s war with Iran',
    'CBS News: Iran war paralyzes oil trade',
    'NPR: Assessing the humanitarian impact of war in Iran',
    'Amnesty International: Urgent call to protect civilians',
    'CFR: The Iran War Is Breaking Global Humanitarian Aid Efforts',
    'OHCHR: Middle East crisis plays out worst fears',
    'Atlantic Council: Twenty questions about the Iran war',
    'Britannica: 2026 Iran Conflict',
    'Wikipedia: 2026 Iran war',
]
for s in en_sources:
    add_bullet(s)

add_sub_heading('한글 출처')
kr_sources = [
    '글로벌이코노믹: 미국-이란 휴전 협상 곧 시작',
    'MBC 뉴스: 장기전으로 미국 고통 주겠다는 이란',
    'MBC 뉴스: 종전 가능성 언급했지만 여전히 불타오르는 중동',
    'MBC 뉴스: 이란, 미국 하르그 섬 공격에 역내 석유시설 반격 경고',
    '문화일보: 이란, 미군 뚫고 호르무즈에 기뢰 설치 시작',
    'KPMG: 자원·물류·AI 3대 축으로 본 미국-이란 전쟁',
    '뉴시스: 이란 전쟁 출구 안갯속 5개 시나리오',
]
for s in kr_sources:
    add_bullet(s)

# 저장
output_path = '/Users/kylechoi/Desktop/Ai works/output/2026-3-14/뉴스크롤링/미국-이란전쟁/미국-이란전쟁_뉴스조사_설교예화.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'Word 파일 생성 완료: {output_path}')
